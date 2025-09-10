import azure.functions as func
from docx import Document
import re
import io
import json

def parse_referral_docx_bytes(b):
    doc = Document(io.BytesIO(b))
    lines = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t: lines.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t: lines.append(t)

    def after(label, strip_extra=True):
        for ln in lines:
            if ln.lower().startswith(label.lower()):
                val = ln.split(":", 1)[1].strip() if ":" in ln else ""
                if strip_extra:
                    val = re.split(r'\t+|\s{2,}', val)[0].strip()
                return val
        return None

    batch_number = None
    image_number = None
    for ln in lines:
        m = re.search(r'Batch Number:\s*(\S+)', ln, re.IGNORECASE)
        if m: batch_number = m.group(1).strip()
        m = re.search(r'Image Number:\s*(\S+)', ln, re.IGNORECASE)
        if m: image_number = m.group(1).strip()

    fields = {
        "FHA CASE NUMBER": after("FHA Case"),
        "BATCH NUMBER": batch_number,
        "IMAGE NUMBER": image_number,
        "SSC REP NAME": after("SSC Rep Name", False),
        "DATE": after("Date"),
        "CERTIFIER": after("Certifier", False),
        "BORROWER (S) NAME": after("Claimant’s Name", False) or after("Claimant's Name", False),
        "MAILING ADDRESS": after("Mailing Address", False),
        "TELEPHONE NUMBER": None,
        "EMAIL ADDRESS": after("Email Address", False),
    }

    for ln in lines:
        if re.search(r"(?i)Telephone", ln):
            m = re.search(r"(\+?\d[\d\-\s()]{6,})", ln)
            if m:
                fields["TELEPHONE NUMBER"] = m.group(1).strip()
                break
    # Capture all lines after "COMMENTS"
    comments = []
    capture = False
    
    for ln in lines:
        if capture:
            if ln.strip() == "":
                break  # Stop at first empty line after comments
            comments.append(ln.strip())
        elif ln.strip().lower() == "comments":
            capture = True
    
    fields["COMMENTS"] = "\n".join(comments) if comments else None
                
    return fields

async def main(req: func.HttpRequest) -> func.HttpResponse:
    try:
        file_bytes = req.get_body()
        if not file_bytes:
            return func.HttpResponse("No file content", status_code=400)
        fields = parse_referral_docx_bytes(file_bytes)
        return func.HttpResponse(json.dumps(fields), mimetype="application/json")
    except Exception as e:
        return func.HttpResponse(str(e), status_code=500)
